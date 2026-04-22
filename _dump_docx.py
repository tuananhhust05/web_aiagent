import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
files = [
    'strategry_feature/Logic Summarization.docx',
    'strategry_feature/Guardrail Logic.docx',
    'strategry_feature/Signal Generation Dynamic way.docx',
    'strategry_feature/Dynamic Startegy Logic.docx',
]
for f in files:
    print(f'\n===== {f} =====')
    d = Document(f)
    for p in d.paragraphs:
        if p.text.strip():
            print(p.text)
    for i, t in enumerate(d.tables):
        print(f'--- TABLE {i} ---')
        for row in t.rows:
            print(' | '.join(c.text.strip() for c in row.cells))
